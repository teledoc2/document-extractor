import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_key(key):
    """Convert snake_case to Title Case and remove underscores"""
    return key.replace('_', ' ').title()

def create_medical_document(json_file_path, output_file_path):
    # Read JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)

    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Medical Form Details', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process OCR contents
    ocr_data = data['ocr_contents']

    # Define sections to process
    sections = ['provider', 'insured', 'patient', 'visitDetails', 'diagnosis', 
               'management', 'services', 'insuranceApproval']

    for section in sections:
        if section in ocr_data:
            section_data = ocr_data[section]
            
            # Skip empty sections
            if not section_data:
                continue

            # Handle special case for services list
            if section == 'services' and section_data:
                # Add section heading
                doc.add_heading('Services', level=1)
                
                for service in section_data:
                    if any(service.values()):  # Check if service has any non-empty values
                        p = doc.add_paragraph()
                        for key, value in service.items():
                            if value and value != "" and value is not False:
                                p.add_run(f"{format_key(key)}: {value}\n")
                continue

            # Add section heading
            if isinstance(section_data, dict) and any(str(v).strip() != "" and v is not False for v in section_data.values()):
                doc.add_heading(format_key(section), level=1)

            # Add key-value pairs
            if isinstance(section_data, dict):
                p = doc.add_paragraph()
                for key, value in section_data.items():
                    # Only add non-empty and non-False values
                    if value and value != "" and value is not False:
                        p.add_run(f"{format_key(key)}: {value}\n")

    # Save the document
    doc.save(output_file_path)

# Execute the function
create_medical_document('outputs/2025-03-19/w_f_abbasia.json', 'outputs/medical_form_abbasia.docx')